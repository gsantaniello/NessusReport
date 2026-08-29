from __future__ import annotations

import csv
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

import nessus_report


ROOT = Path(__file__).resolve().parents[1]
SAMPLE = ROOT / "samples" / "sample_scan.nessus"
TEMPLATE = ROOT / "Nessus_Report_Template.docx"


class NessusReportTests(unittest.TestCase):
    def test_parse_sample(self) -> None:
        findings = nessus_report.load_findings([SAMPLE])
        self.assertEqual(len(findings), 5)
        self.assertEqual({finding.host_ip for finding in findings}, {"192.0.2.10", "192.0.2.20"})
        self.assertEqual([finding.severity for finding in findings], [4, 2, 2, 1, 0])
        self.assertEqual(sum(finding.plugin_id == 100002 for finding in findings), 2)

    def test_csv_generation(self) -> None:
        findings = nessus_report.load_findings([SAMPLE])
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "report.csv"
            nessus_report.write_csv(findings, output)
            with output.open("r", encoding="utf-8-sig", newline="") as handle:
                rows = list(csv.DictReader(handle, delimiter="|"))
        self.assertEqual(len(rows), 5)
        self.assertEqual(rows[0]["Severity Label"], "Critical")
        self.assertEqual(rows[-1]["Severity Label"], "Info")
        self.assertEqual(rows[0]["CVEs"], "CVE-2099-0001")

    def test_docx_generation(self) -> None:
        findings = nessus_report.load_findings([SAMPLE])
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "report.docx"
            nessus_report.generate_docx(findings, TEMPLATE, output, [SAMPLE])
            self.assertTrue(zipfile.is_zipfile(output))
            document = Document(output)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        combined = f"{text}\n{table_text}"
        self.assertIn("Nessus Vulnerability Report", combined)
        self.assertIn("Unsupported web server version", combined)
        self.assertIn("TLS weak cipher suites enabled", combined)
        self.assertIn("CRITICAL", combined)
        self.assertIn("INFO", combined)

    def test_rejects_doctype_and_entities(self) -> None:
        unsafe_xml = b'''<?xml version="1.0"?>
<!DOCTYPE sample [<!ENTITY external SYSTEM "file:///etc/passwd">]>
<NessusClientData_v2><Report name="unsafe">&external;</Report></NessusClientData_v2>
'''
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "unsafe.nessus"
            path.write_bytes(unsafe_xml)
            with self.assertRaises(nessus_report.NessusReportError):
                list(nessus_report.parse_nessus(path))


if __name__ == "__main__":
    unittest.main()
