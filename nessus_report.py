#!/usr/bin/env python3
"""Convert Nessus v2 XML reports to a pipe-delimited CSV and a DOCX report."""

from __future__ import annotations

import argparse
import csv
import logging
import os
import re
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterator, Sequence

from lxml import etree


__version__ = "1.0.0"

LOGGER = logging.getLogger("nessus_report")

SEVERITY_LABELS = {
    0: "Info",
    1: "Low",
    2: "Medium",
    3: "High",
    4: "Critical",
}

SEVERITY_COLORS = {
    0: "2E7D32",
    1: "0288D1",
    2: "ED9C18",
    3: "D32F2F",
    4: "7030A0",
}

CSV_FIELDS = [
    "IP",
    "Target",
    "FQDN",
    "Asset Name",
    "Asset Operating System",
    "Port",
    "Protocol",
    "Service",
    "Severity",
    "Severity Label",
    "Risk Factor",
    "Plugin ID",
    "Vulnerability Title",
    "CVSS2 Base Score",
    "CVSS2 Vector",
    "CVSS3 Base Score",
    "CVSS3 Vector",
    "CVEs",
    "Description",
    "Mitigation",
    "Output",
]


class NessusReportError(RuntimeError):
    """Raised for controlled input or output errors."""


@dataclass(frozen=True, slots=True)
class Finding:
    source: str
    target: str
    host_ip: str
    host_fqdn: str
    netbios_name: str
    operating_system: str
    mac_address: str
    port: int
    protocol: str
    service: str
    severity: int
    plugin_id: int
    plugin_name: str
    plugin_family: str
    risk_factor: str
    synopsis: str
    description: str
    solution: str
    plugin_output: str
    cvss2_base_score: str
    cvss2_vector: str
    cvss2_temporal_score: str
    cvss2_temporal_vector: str
    cvss3_base_score: str
    cvss3_vector: str
    cvss3_temporal_score: str
    cvss3_temporal_vector: str
    cves: tuple[str, ...]
    bids: tuple[str, ...]
    xrefs: tuple[str, ...]

    @property
    def severity_label(self) -> str:
        return SEVERITY_LABELS.get(self.severity, f"Unknown ({self.severity})")

    @property
    def endpoint(self) -> str:
        protocol = self.protocol.upper() or "N/A"
        return f"{self.port}/{protocol}"

    @property
    def asset_label(self) -> str:
        return self.host_fqdn or self.target or self.host_ip or "Unknown asset"


def _normalise_text(value: str | None, *, multiline: bool = False) -> str:
    if not value:
        return ""
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    if multiline:
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.split("\n")]
        return "\n".join(line for line in lines if line)
    return re.sub(r"\s+", " ", value).strip()


def _safe_int(value: str | None, default: int = 0) -> int:
    try:
        return int(value or default)
    except (TypeError, ValueError):
        return default


def _first(values: dict[str, list[str]], key: str) -> str:
    entries = values.get(key, [])
    return entries[0] if entries else ""


def _strip_cvss_prefix(vector: str) -> str:
    vector = _normalise_text(vector)
    if "#" in vector:
        return vector.split("#", 1)[1]
    if vector.upper().startswith("CVSS:3.0/") or vector.upper().startswith("CVSS:3.1/"):
        return vector.split("/", 1)[1]
    return vector


def _assert_safe_xml(path: Path) -> None:
    try:
        with path.open("rb") as handle:
            prefix = handle.read(16_384).upper()
    except OSError as exc:
        raise NessusReportError(f"Cannot read input file {path}: {exc}") from exc
    if b"<!DOCTYPE" in prefix or b"<!ENTITY" in prefix:
        raise NessusReportError(
            f"Unsafe XML declaration detected in {path}; DTD and entity declarations are not accepted."
        )


def parse_nessus(path: Path) -> Iterator[Finding]:
    """Yield findings from a Nessus v2 XML file using bounded-memory parsing."""

    if not path.is_file():
        raise NessusReportError(f"Input file does not exist: {path}")
    _assert_safe_xml(path)

    host = {
        "target": "",
        "host-ip": "",
        "host-fqdn": "",
        "netbios-name": "",
        "operating-system": "",
        "mac-address": "",
    }

    try:
        context = etree.iterparse(
            str(path),
            events=("start", "end"),
            load_dtd=False,
            no_network=True,
            resolve_entities=False,
            huge_tree=False,
            recover=False,
        )

        for event, elem in context:
            if event == "start" and elem.tag == "ReportHost":
                host = {
                    "target": _normalise_text(elem.get("name")),
                    "host-ip": "",
                    "host-fqdn": "",
                    "netbios-name": "",
                    "operating-system": "",
                    "mac-address": "",
                }
                continue

            if event != "end":
                continue

            if elem.tag == "tag":
                name = elem.get("name", "")
                if name in host:
                    host[name] = _normalise_text(elem.text, multiline=name == "operating-system")
                continue

            if elem.tag == "ReportItem":
                child_values: dict[str, list[str]] = defaultdict(list)
                for child in elem:
                    child_values[child.tag].append(
                        _normalise_text(child.text, multiline=child.tag in {"plugin_output", "description", "solution"})
                    )

                severity = _safe_int(elem.get("severity"), 0)
                finding = Finding(
                    source=path.name,
                    target=host["target"],
                    host_ip=host["host-ip"],
                    host_fqdn=host["host-fqdn"],
                    netbios_name=host["netbios-name"],
                    operating_system=host["operating-system"],
                    mac_address=host["mac-address"],
                    port=_safe_int(elem.get("port"), 0),
                    protocol=_normalise_text(elem.get("protocol")).upper(),
                    service=_normalise_text(elem.get("svc_name")),
                    severity=severity,
                    plugin_id=_safe_int(elem.get("pluginID"), 0),
                    plugin_name=_normalise_text(elem.get("pluginName")) or "Unnamed Nessus finding",
                    plugin_family=_normalise_text(elem.get("pluginFamily")),
                    risk_factor=_first(child_values, "risk_factor") or SEVERITY_LABELS.get(severity, "Unknown"),
                    synopsis=_first(child_values, "synopsis"),
                    description=_first(child_values, "description"),
                    solution=_first(child_values, "solution"),
                    plugin_output=_first(child_values, "plugin_output"),
                    cvss2_base_score=_first(child_values, "cvss_base_score"),
                    cvss2_vector=_strip_cvss_prefix(_first(child_values, "cvss_vector")),
                    cvss2_temporal_score=_first(child_values, "cvss_temporal_score"),
                    cvss2_temporal_vector=_strip_cvss_prefix(_first(child_values, "cvss_temporal_vector")),
                    cvss3_base_score=_first(child_values, "cvss3_base_score"),
                    cvss3_vector=_strip_cvss_prefix(_first(child_values, "cvss3_vector")),
                    cvss3_temporal_score=_first(child_values, "cvss3_temporal_score"),
                    cvss3_temporal_vector=_strip_cvss_prefix(_first(child_values, "cvss3_temporal_vector")),
                    cves=tuple(v for v in child_values.get("cve", []) if v),
                    bids=tuple(v for v in child_values.get("bid", []) if v),
                    xrefs=tuple(v for v in child_values.get("xref", []) if v),
                )
                yield finding
                elem.clear()
                while elem.getprevious() is not None:
                    del elem.getparent()[0]
                continue

            if elem.tag == "ReportHost":
                elem.clear()
    except (etree.XMLSyntaxError, OSError) as exc:
        raise NessusReportError(f"Cannot parse Nessus XML {path}: {exc}") from exc


def load_findings(paths: Sequence[Path]) -> list[Finding]:
    findings: list[Finding] = []
    seen: set[Finding] = set()
    for path in paths:
        for finding in parse_nessus(path):
            if finding not in seen:
                findings.append(finding)
                seen.add(finding)
    findings.sort(key=lambda item: (-item.severity, item.plugin_name.lower(), item.host_ip, item.port))
    return findings


def finding_to_csv_row(finding: Finding) -> dict[str, str | int]:
    return {
        "IP": finding.host_ip,
        "Target": finding.target,
        "FQDN": finding.host_fqdn,
        "Asset Name": finding.netbios_name,
        "Asset Operating System": finding.operating_system,
        "Port": finding.port,
        "Protocol": finding.protocol,
        "Service": finding.service,
        "Severity": finding.severity,
        "Severity Label": finding.severity_label,
        "Risk Factor": finding.risk_factor,
        "Plugin ID": finding.plugin_id,
        "Vulnerability Title": finding.plugin_name,
        "CVSS2 Base Score": finding.cvss2_base_score,
        "CVSS2 Vector": finding.cvss2_vector,
        "CVSS3 Base Score": finding.cvss3_base_score,
        "CVSS3 Vector": finding.cvss3_vector,
        "CVEs": ", ".join(finding.cves),
        "Description": finding.description,
        "Mitigation": finding.solution,
        "Output": finding.plugin_output,
    }


def write_csv(findings: Sequence[Finding], destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_name(f".{destination.name}.tmp")
    try:
        with temporary.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(
                handle,
                fieldnames=CSV_FIELDS,
                delimiter="|",
                quotechar='"',
                quoting=csv.QUOTE_MINIMAL,
                lineterminator="\n",
            )
            writer.writeheader()
            writer.writerows(finding_to_csv_row(item) for item in findings)
        os.replace(temporary, destination)
    except OSError as exc:
        if temporary.exists():
            temporary.unlink()
        raise NessusReportError(f"Cannot write CSV {destination}: {exc}") from exc


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 9.5) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text or "—")
    run.bold = bold
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def _set_cell_width(cell, width_dxa: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: Sequence[int], *, indent_dxa: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    for tag, value in (("w:tblW", total), ("w:tblInd", indent_dxa)):
        element = tbl_pr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tbl_pr.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            _set_cell_width(cell, widths_dxa[min(index, len(widths_dxa) - 1)])


def _set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def _group_findings(findings: Sequence[Finding]) -> list[list[Finding]]:
    groups: dict[tuple[int, str, int, str, str, int], list[Finding]] = defaultdict(list)
    for finding in findings:
        key = (
            finding.plugin_id,
            finding.plugin_name,
            finding.port,
            finding.protocol,
            finding.service,
            finding.severity,
        )
        groups[key].append(finding)
    return sorted(
        groups.values(),
        key=lambda group: (-group[0].severity, group[0].plugin_name.lower(), group[0].plugin_id),
    )


def _add_label_value(document, label: str, value: str) -> None:
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    paragraph.add_run(value or "Not provided")


def generate_docx(
    findings: Sequence[Finding],
    template_path: Path,
    destination: Path,
    source_paths: Sequence[Path],
) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise NessusReportError("python-docx is required to generate DOCX reports") from exc

    if not template_path.is_file():
        raise NessusReportError(f"DOCX template does not exist: {template_path}")

    try:
        document = Document(str(template_path))
    except (OSError, ValueError) as exc:
        raise NessusReportError(f"Cannot open DOCX template {template_path}: {exc}") from exc

    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    core = document.core_properties
    core.title = "Nessus Vulnerability Report"
    core.subject = "Generated locally from Nessus v2 XML input"
    core.author = ""
    core.last_modified_by = ""
    core.comments = "Generated by NessusReport 1.0.0"

    title = document.add_paragraph(style="Report Title")
    title.add_run("Nessus Vulnerability Report")
    subtitle = document.add_paragraph(style="Report Subtitle")
    subtitle.add_run("Structured findings generated locally from Nessus v2 XML")

    metadata = document.add_table(rows=4, cols=2)
    metadata.style = "Table Grid"
    _set_table_geometry(metadata, [2500, 6860])
    metadata_rows = [
        ("Generated", datetime.now().astimezone().strftime("%Y-%m-%d %H:%M %Z")),
        ("Source files", ", ".join(path.name for path in source_paths)),
        ("Assets", str(len({(item.host_ip, item.target) for item in findings}))),
        ("Findings", f"{len(findings)} instances / {len(_group_findings(findings))} grouped findings"),
    ]
    for row, (label, value) in zip(metadata.rows, metadata_rows):
        _shade_cell(row.cells[0], "E8EEF5")
        _set_cell_text(row.cells[0], label, bold=True, color="1F4D78")
        _set_cell_text(row.cells[1], value)
        for cell in row.cells:
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_paragraph()
    heading = document.add_paragraph("Executive summary", style="Heading 1")
    _set_keep_with_next(heading)
    summary = Counter(item.severity for item in findings)

    summary_table = document.add_table(rows=2, cols=6)
    summary_table.style = "Table Grid"
    _set_table_geometry(summary_table, [1560] * 6)
    headers = ["Total", "Critical", "High", "Medium", "Low", "Info"]
    values = [len(findings), summary[4], summary[3], summary[2], summary[1], summary[0]]
    fills = ["1F4D78", SEVERITY_COLORS[4], SEVERITY_COLORS[3], SEVERITY_COLORS[2], SEVERITY_COLORS[1], SEVERITY_COLORS[0]]
    for index, header in enumerate(headers):
        _shade_cell(summary_table.rows[0].cells[index], fills[index])
        _set_cell_text(summary_table.rows[0].cells[index], header, bold=True, color="FFFFFF")
        _set_cell_text(summary_table.rows[1].cells[index], str(values[index]), bold=True, size=12)
        for cell in (summary_table.rows[0].cells[index], summary_table.rows[1].cells[index]):
            _set_cell_margins(cell, top=120, bottom=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    findings_section = document.add_section(WD_SECTION.NEW_PAGE)
    findings_section.header.is_linked_to_previous = True
    findings_section.footer.is_linked_to_previous = True
    heading = document.add_paragraph("Findings", style="Heading 1")
    _set_keep_with_next(heading)

    grouped = _group_findings(findings)
    for group_index, group in enumerate(grouped, start=1):
        representative = group[0]
        banner = document.add_table(rows=1, cols=2)
        banner.style = "Table Grid"
        _set_table_geometry(banner, [7560, 1800])
        _shade_cell(banner.cell(0, 0), "E8EEF5")
        _shade_cell(banner.cell(0, 1), SEVERITY_COLORS.get(representative.severity, "666666"))
        _set_cell_text(
            banner.cell(0, 0),
            f"{group_index}. {representative.plugin_name}",
            bold=True,
            color="0B2545",
            size=12,
        )
        _set_cell_text(
            banner.cell(0, 1),
            representative.severity_label.upper(),
            bold=True,
            color="FFFFFF",
            size=10,
        )
        banner.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in banner.rows[0].cells:
            _set_cell_margins(cell, top=130, bottom=130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        facts = document.add_table(rows=4, cols=4)
        facts.style = "Table Grid"
        _set_table_geometry(facts, [1500, 3180, 1500, 3180])
        assets = sorted({item.asset_label for item in group})
        cves = sorted({cve for item in group for cve in item.cves})
        cvss = representative.cvss3_base_score or representative.cvss2_base_score or "Not provided"
        fact_values = [
            ("Plugin ID", str(representative.plugin_id), "CVSS", cvss),
            ("Service", representative.service or "Not provided", "Endpoint", representative.endpoint),
            ("Assets", ", ".join(assets), "Instances", str(len(group))),
            ("CVEs", ", ".join(cves) or "Not provided", "Family", representative.plugin_family or "Not provided"),
        ]
        for row, values_row in zip(facts.rows, fact_values):
            for index, value in enumerate(values_row):
                if index % 2 == 0:
                    _shade_cell(row.cells[index], "F2F4F7")
                    _set_cell_text(row.cells[index], value, bold=True, color="1F4D78")
                else:
                    _set_cell_text(row.cells[index], value)
                _set_cell_margins(row.cells[index])
                row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        document.add_paragraph()
        if representative.synopsis:
            _add_label_value(document, "Synopsis", representative.synopsis)
        _add_label_value(document, "Description", representative.description)
        _add_label_value(document, "Remediation", representative.solution)

        evidence_heading = document.add_paragraph("Evidence", style="Heading 3")
        _set_keep_with_next(evidence_heading)
        for finding in group:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            asset_run = p.add_run(f"{finding.asset_label} ({finding.host_ip or 'IP unavailable'})\n")
            asset_run.bold = True
            output_run = p.add_run(finding.plugin_output or "No plugin output recorded")
            output_run.font.name = "Courier New"
            output_run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Courier New")
            output_run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Courier New")
            output_run.font.size = Pt(8.5)
            output_run.font.color.rgb = RGBColor.from_string("333333")

        if group_index < len(grouped):
            document.add_page_break()

    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_name(f".{destination.name}.tmp")
    try:
        document.save(str(temporary))
        with zipfile.ZipFile(temporary, "r") as archive:
            bad_member = archive.testzip()
            if bad_member:
                raise NessusReportError(f"Generated DOCX contains a corrupt member: {bad_member}")
        os.replace(temporary, destination)
    except (OSError, zipfile.BadZipFile) as exc:
        if temporary.exists():
            temporary.unlink()
        raise NessusReportError(f"Cannot write DOCX {destination}: {exc}") from exc


def _expand_inputs(raw_inputs: Sequence[str]) -> list[Path]:
    expanded: list[Path] = []
    for raw in raw_inputs:
        for value in raw.split(","):
            value = value.strip()
            if value:
                expanded.append(Path(value).expanduser().resolve())
    if not expanded:
        raise NessusReportError("At least one .nessus input file is required")
    return expanded


def build_argument_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="nessus-report",
        description="Convert one or more Nessus v2 XML files to CSV and DOCX reports.",
    )
    parser.add_argument("inputs", nargs="+", help="One or more .nessus files; comma-separated input is also supported")
    parser.add_argument(
        "-t",
        "--template",
        type=Path,
        default=Path(__file__).with_name("Nessus_Report_Template.docx"),
        help="DOCX template (default: Nessus_Report_Template.docx next to the script)",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        required=True,
        help="Output base path without extension",
    )
    output_group = parser.add_mutually_exclusive_group()
    output_group.add_argument("--csv-only", action="store_true", help="Generate only the CSV report")
    output_group.add_argument("--docx-only", action="store_true", help="Generate only the DOCX report")
    parser.add_argument("--verbose", action="store_true", help="Enable diagnostic logging")
    parser.add_argument("--version", action="version", version=f"%(prog)s {__version__}")
    return parser


def run(argv: Sequence[str] | None = None) -> int:
    parser = build_argument_parser()
    args = parser.parse_args(argv)
    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(levelname)s: %(message)s",
    )

    try:
        inputs = _expand_inputs(args.inputs)
        findings = load_findings(inputs)
        if not findings:
            raise NessusReportError("No ReportItem findings were found in the supplied Nessus files")

        output_base = args.output.expanduser().resolve()
        csv_path = output_base.with_suffix(".csv")
        docx_path = output_base.with_suffix(".docx")

        if not args.docx_only:
            write_csv(findings, csv_path)
            LOGGER.info("CSV written to %s", csv_path)
        if not args.csv_only:
            generate_docx(findings, args.template.expanduser().resolve(), docx_path, inputs)
            LOGGER.info("DOCX written to %s", docx_path)

        counts = Counter(item.severity_label for item in findings)
        summary = ", ".join(f"{label}={counts[label]}" for label in ("Critical", "High", "Medium", "Low", "Info"))
        print(f"Processed {len(findings)} findings across {len({item.host_ip for item in findings})} assets ({summary}).")
        return 0
    except NessusReportError as exc:
        LOGGER.error("%s", exc)
        return 2


def main() -> None:
    raise SystemExit(run())


if __name__ == "__main__":
    main()
