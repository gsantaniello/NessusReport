#!/usr/bin/env python3
"""Build the deterministic DOCX style template used by nessus_report.py."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor.from_string("2E74B5")
DARK_BLUE = RGBColor.from_string("1F4D78")
INK = RGBColor.from_string("0B2545")
MUTED = RGBColor.from_string("667085")


def set_font(style, name: str, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def build_template(destination: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    set_font(normal, "Calibri", 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        set_font(style, "Calibri", size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    report_title = styles.add_style("Report Title", 1)
    set_font(report_title, "Calibri", 26, INK, True)
    report_title.paragraph_format.space_before = Pt(0)
    report_title.paragraph_format.space_after = Pt(6)
    report_title.paragraph_format.keep_with_next = True

    report_subtitle = styles.add_style("Report Subtitle", 1)
    set_font(report_subtitle, "Calibri", 12, MUTED, False)
    report_subtitle.paragraph_format.space_before = Pt(0)
    report_subtitle.paragraph_format.space_after = Pt(20)
    report_subtitle.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    header_run = header_p.add_run("NESSUS REPORT  •  GENERATED LOCALLY")
    header_run.font.name = "Calibri"
    header_run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    header_run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    header_run.font.size = Pt(8)
    header_run.font.bold = True
    header_run.font.color.rgb = MUTED

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    footer_table.autofit = False
    set_cell_width(footer_table.cell(0, 0), 7200)
    set_cell_width(footer_table.cell(0, 1), 2160)
    footer_table.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    footer_table.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left = footer_table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    left_run = left.add_run("Nessus Vulnerability Report")
    left_run.font.name = "Calibri"
    left_run.font.size = Pt(8)
    left_run.font.color.rgb = MUTED
    right = footer_table.cell(0, 1).paragraphs[0]
    right.paragraph_format.space_after = Pt(0)
    add_page_field(right)
    for run in right.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED

    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    core = document.core_properties
    core.title = "Nessus Report Template"
    core.subject = "Style template for NessusReport 1.0.0"
    core.author = ""
    core.last_modified_by = ""
    core.comments = "Generated deterministically by tools/build_template.py"

    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    build_template(args.output.resolve())


if __name__ == "__main__":
    main()
